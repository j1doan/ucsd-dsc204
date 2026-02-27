#!/usr/bin/env python3
"""Generate HW3 Word report. Uses all plots, meets assignment requirements."""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUT_DIR = os.path.dirname(os.path.abspath(__file__))
FIG_W = Inches(3.4)

def add_section(doc, title, body, fig_path=None):
    p = doc.add_paragraph()
    p.add_run(title).bold = True
    p.space_after = Pt(3)
    doc.add_paragraph(body).space_after = Pt(3)
    if fig_path and os.path.exists(fig_path):
        doc.add_picture(fig_path, width=FIG_W)
        doc.add_paragraph().space_after = Pt(4)

def add_figs(doc, fnames):
    for fname in fnames:
        path = os.path.join(OUT_DIR, fname)
        if os.path.exists(path):
            doc.add_picture(path, width=FIG_W)
            doc.add_paragraph().space_after = Pt(3)

def main():
    doc = Document()
    doc.styles['Normal'].font.size = Pt(10)

    # Title
    p = doc.add_paragraph()
    run = p.add_run("Taxi Fare Prediction with a GAM: Summary Report")
    run.bold = True
    run.font.size = Pt(13)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # 1. Goal and Setup
    add_section(doc,
        "1. Goal and Setup",
        "We built a GAM to predict NYC yellow taxi fare from non-fare trip features. The data come from January 2021 "
        "yellow taxi Parquet files: we used 100k rows with an 80/20 train/validation split. Predictors include trip "
        "distance, trip duration (derived from pickup and dropoff times), hour of day, day of week, passenger count, "
        "and HW2 PC1/PC2 scores joined by pickup zone. No fare-related fields such as tolls or tips were used."
    )

    # 2. Model and Evaluation
    add_section(doc,
        "2. Model and Evaluation",
        "We fit a pygam LinearGAM with smooth spline terms for distance, duration, hour, and day of week, and a linear "
        "term for passenger count. On the validation set we get RMSE=$2.16, MAE=$0.75, and R²=0.924. The scatter plot "
        "below shows actual vs predicted fares against the y=x reference line. Most points sit close to the line, which "
        "suggests the model fits well overall. There is more scatter at higher fare levels, so predictions for "
        "expensive trips tend to be less precise."
        ,
        os.path.join(OUT_DIR, "actual_vs_predicted.png")
    )

    # 3. Partial Dependence
    add_section(doc,
        "3. Partial Dependence (Interpretation)",
        "The partial dependence plots show how each predictor relates to fare when the others are held at a reference. "
        "Distance and duration both push fare up in a roughly linear way, which matches meter-based pricing. Hour of "
        "day has a modest dip around midday and a slight evening peak, reflecting when demand and fares tend to rise. "
        "Day of week and passenger count have smaller effects. The 95% confidence bands are fairly narrow, so these "
        "patterns look stable."
        ,
        os.path.join(OUT_DIR, "partial_dependence.png")
    )

    # 4. Bootstrap CI [EC]
    add_section(doc,
        "4. Bootstrap CI Comparison [EC]",
        "We compared bootstrap 95% confidence intervals (50 resamples) with pygam’s built-in intervals for distance, "
        "duration, and hour. The two methods agree closely, so the model’s uncertainty estimates appear reasonable."
        ,
        os.path.join(OUT_DIR, "bootstrap_ci_comparison.png")
    )

    # 5. Fare breakdown [EC]
    add_section(doc,
        "5. Per-Term Fare Breakdown [EC]",
        "The bar chart below breaks down how much each predictor contributes to the average predicted fare. Distance "
        "and duration dominate, which makes sense since they drive meter charges. Hour, day of week, and passenger "
        "count add smaller contributions."
        ,
        os.path.join(OUT_DIR, "fare_component_breakdown.png")
    )

    # 6. EXTRA 1 — Seaborn EDA
    p = doc.add_paragraph()
    p.add_run("6. EXTRA 1 — Seaborn EDA").bold = True
    p.space_after = Pt(3)
    doc.add_paragraph(
        "The first plot shows the fare distribution: it’s right-skewed, and the log scale helps normalize it. The violin "
        "by hour of day shows that fares vary across the day, with different shapes in the morning, midday, and evening. "
        "The third plot compares fare by day of week; weekends and weekdays show some differences. The correlation "
        "heatmap indicates that distance and duration are strongly correlated with fare and with each other, while "
        "hour and day have weaker associations."
    ).space_after = Pt(3)
    add_figs(doc, [
        "extra_fare_distribution.png",
        "extra_fare_by_hour_violin.png",
        "extra_fare_by_dow.png",
        "extra_correlation_heatmap.png",
    ])

    # 7. EXTRA 2 — Region Analysis
    p = doc.add_paragraph()
    p.add_run("7. EXTRA 2 — Region Analysis").bold = True
    p.space_after = Pt(3)
    doc.add_paragraph(
        "We assigned pickup zones to NYC boroughs using zone centroids. The first plot shows fare distributions by "
        "pickup region; Manhattan and other boroughs differ in typical fare levels. The second plot shows the GAM "
        "distance effect fitted separately by region. The slope of fare vs distance can vary by borough, suggesting "
        "that the relationship between distance and fare is not uniform across the city."
    ).space_after = Pt(3)
    add_figs(doc, [
        "extra_fare_by_region.png",
        "extra_gam_by_region.png",
    ])

    # 8. EXTRA 3 — Multi-year & Multi-service
    p = doc.add_paragraph()
    p.add_run("8. EXTRA 3 — Multi-year & Multi-service").bold = True
    p.space_after = Pt(3)
    doc.add_paragraph(
        "We sampled from Yellow taxi in 2015, 2021, and 2022, plus Green taxi in 2021. The first plot compares fare "
        "distributions across services and years; levels and shapes differ, reflecting changes in rates and ridership. "
        "The second plot shows distance vs fare by service; the slope and spread vary. The third plot overlays the GAM "
        "distance partial-dependence curves for each service. The curves differ in level and shape, which suggests "
        "that fare structures have changed over time and between yellow and green cabs."
    ).space_after = Pt(3)
    add_figs(doc, [
        "extra_fare_by_service.png",
        "extra_distance_vs_fare_by_service.png",
        "extra_gam_distance_by_service.png",
    ])

    # 9. Limitations
    p = doc.add_paragraph()
    p.add_run("9. Limitations").bold = True
    p.space_after = Pt(3)
    doc.add_paragraph(
        "The model excludes tolls, MTA tax, congestion surcharges, and airport fees, so trips through toll zones will "
        "be under-predicted. We do not use fine-grained GPS; route geometry and traffic are not captured. Training is "
        "on January 2021 only, so seasonal and post-pandemic effects are not reflected."
    ).space_after = Pt(6)

    out_path = os.path.join(OUT_DIR, "HW3_GAM_Fare_Prediction_Report.docx")
    doc.save(out_path)
    print(f"Report saved to {out_path}")

if __name__ == "__main__":
    main()
